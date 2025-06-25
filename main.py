#!/usr/bin/env python3

import os
import shutil
import json
import magic
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from datetime import datetime
import hashlib
import logging
from dataclasses import dataclass
from enum import Enum
import time

from PIL import Image
import docx
import PyPDF2
import pandas as pd
from openpyxl import load_workbook
import zipfile
import subprocess

from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

from colorama import init, Fore, Style, Back
init()

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('system.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

try:
    from ai_analyzer import FreeAIAnalyzer, check_dependencies
    AI_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ AI module not available: {e}")
    AI_AVAILABLE = False

class DocumentType(Enum):
    DOCUMENTS = "documents"
    IMAGES = "images"
    VIDEOS = "videos"
    AUDIO = "audio"
    CODE = "code"
    DATA = "data"
    PRESENTATIONS = "presentations"
    COMPRESSED_FILES = "compressed"
    EXECUTABLES = "executables"
    OTHERS = "others"

@dataclass
class FileMetadata:
    path: Path
    name: str
    extension: str
    size: int
    created_date: datetime
    modified_date: datetime
    mime_type: str
    content_summary: str
    category: DocumentType
    subcategory: str
    hash_md5: str

class FileAnalyzer:
    def __init__(self):
        self.supported_extensions = {
            '.pdf', '.docx', '.doc', '.txt', '.rtf', '.odt',
            '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.svg', '.webp',
            '.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm',
            '.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma',
            '.py', '.js', '.html', '.css', '.java', '.cpp', '.c', '.php', '.rb', '.go',
            '.xlsx', '.xls', '.csv', '.json', '.xml', '.sql',
            '.pptx', '.ppt', '.odp',
            '.zip', '.rar', '.7z', '.tar', '.gz',
            '.exe', '.msi', '.deb', '.rpm', '.dmg'
        }
        
        self.ai_analyzer = None
        if AI_AVAILABLE:
            try:
                self.ai_analyzer = FreeAIAnalyzer()
                logger.info("Free AI analyzer initialized")
            except Exception as e:
                logger.warning(f"Error initializing free AI: {e}")
                self.ai_analyzer = None
    
    def get_file_hash(self, file_path: Path) -> str:
        hash_md5 = hashlib.md5()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            logger.error(f"Error calculating hash for {file_path}: {e}")
            return ""
    
    def analyze_file(self, file_path: Path) -> FileMetadata:
        try:
            stat = file_path.stat()
            name = file_path.name
            extension = file_path.suffix.lower()
            size = stat.st_size
            created_date = datetime.fromtimestamp(stat.st_ctime)
            modified_date = datetime.fromtimestamp(stat.st_mtime)
            
            mime_type = magic.from_file(str(file_path), mime=True)
            file_hash = self.get_file_hash(file_path)
            content_summary = self._analyze_content(file_path, extension, mime_type)
            category, subcategory = self._categorize_file(file_path, extension, mime_type, content_summary)
            
            return FileMetadata(
                path=file_path,
                name=name,
                extension=extension,
                size=size,
                created_date=created_date,
                modified_date=modified_date,
                mime_type=mime_type,
                content_summary=content_summary,
                category=category,
                subcategory=subcategory,
                hash_md5=file_hash
            )
            
        except Exception as e:
            logger.error(f"Error analyzing file {file_path}: {e}")
            return None
    
    def _analyze_content(self, file_path: Path, extension: str, mime_type: str) -> str:
        try:
            if extension in ['.txt', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv']:
                return self._analyze_text_file(file_path)
            elif extension == '.pdf':
                return self._analyze_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                return self._analyze_word_doc(file_path)
            elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                return self._analyze_image(file_path)
            elif extension in ['.xlsx', '.xls']:
                return self._analyze_excel(file_path)
            elif extension in ['.zip', '.rar', '.7z']:
                return self._analyze_compressed_file(file_path)
            else:
                return f"{extension} file - Type: {mime_type}"
        except Exception as e:
            logger.error(f"Error analyzing content of {file_path}: {e}")
            return f"Error analyzing content: {str(e)}"
    
    def _analyze_text_file(self, file_path: Path) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(1000)
                lines = len(content.split('\n'))
                words = len(content.split())
                return f"Text file - {lines} lines, {words} words"
        except Exception as e:
            return f"Error reading text file: {e}"
    
    def _analyze_pdf(self, file_path: Path) -> str:
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                pages = len(reader.pages)
                if pages > 0:
                    first_page = reader.pages[0].extract_text()[:200]
                    return f"PDF - {pages} pages. Content: {first_page}..."
                return f"PDF - {pages} pages"
        except Exception as e:
            return f"Error reading PDF: {e}"
    
    def _analyze_word_doc(self, file_path: Path) -> str:
        try:
            doc = docx.Document(file_path)
            paragraphs = len(doc.paragraphs)
            text_preview = ""
            for para in doc.paragraphs[:3]:
                text_preview += para.text + " "
            return f"Word document - {paragraphs} paragraphs. Content: {text_preview[:200]}..."
        except Exception as e:
            return f"Error reading Word document: {e}"
    
    def _analyze_image(self, file_path: Path) -> str:
        try:
            with Image.open(file_path) as img:
                width, height = img.size
                mode = img.mode
                format_img = img.format
                return f"{format_img} image - {width}x{height} pixels, mode {mode}"
        except Exception as e:
            return f"Error analyzing image: {e}"
    
    def _analyze_excel(self, file_path: Path) -> str:
        try:
            df = pd.read_excel(file_path, nrows=0)
            sheets = pd.ExcelFile(file_path).sheet_names
            columns = len(df.columns)
            return f"Excel - {len(sheets)} sheets, {columns} columns in first sheet"
        except Exception as e:
            return f"Error analyzing Excel: {e}"
    
    def _analyze_compressed_file(self, file_path: Path) -> str:
        try:
            if file_path.suffix.lower() == '.zip':
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    files = len(zip_file.namelist())
                    return f"ZIP file - {files} compressed files"
            else:
                return f"Compressed file {file_path.suffix}"
        except Exception as e:
            return f"Error analyzing compressed file: {e}"
    
    def _categorize_file(self, file_path: Path, extension: str, mime_type: str, content: str) -> Tuple[DocumentType, str]:
        category_map = {
            '.pdf': (DocumentType.DOCUMENTS, 'pdf'),
            '.docx': (DocumentType.DOCUMENTS, 'word'),
            '.doc': (DocumentType.DOCUMENTS, 'word'),
            '.txt': (DocumentType.DOCUMENTS, 'text'),
            '.rtf': (DocumentType.DOCUMENTS, 'rich_text'),
            '.odt': (DocumentType.DOCUMENTS, 'libreoffice'),
            
            '.jpg': (DocumentType.IMAGES, 'jpeg'),
            '.jpeg': (DocumentType.IMAGES, 'jpeg'),
            '.png': (DocumentType.IMAGES, 'png'),
            '.gif': (DocumentType.IMAGES, 'gif'),
            '.bmp': (DocumentType.IMAGES, 'bitmap'),
            '.tiff': (DocumentType.IMAGES, 'tiff'),
            '.svg': (DocumentType.IMAGES, 'vector'),
            '.webp': (DocumentType.IMAGES, 'webp'),
            
            '.mp4': (DocumentType.VIDEOS, 'mp4'),
            '.avi': (DocumentType.VIDEOS, 'avi'),
            '.mkv': (DocumentType.VIDEOS, 'mkv'),
            '.mov': (DocumentType.VIDEOS, 'quicktime'),
            '.wmv': (DocumentType.VIDEOS, 'windows_media'),
            '.flv': (DocumentType.VIDEOS, 'flash'),
            '.webm': (DocumentType.VIDEOS, 'webm'),
            
            '.mp3': (DocumentType.AUDIO, 'mp3'),
            '.wav': (DocumentType.AUDIO, 'wav'),
            '.flac': (DocumentType.AUDIO, 'flac'),
            '.aac': (DocumentType.AUDIO, 'aac'),
            '.ogg': (DocumentType.AUDIO, 'ogg'),
            '.wma': (DocumentType.AUDIO, 'windows_media'),
            
            '.py': (DocumentType.CODE, 'python'),
            '.js': (DocumentType.CODE, 'javascript'),
            '.html': (DocumentType.CODE, 'html'),
            '.css': (DocumentType.CODE, 'css'),
            '.java': (DocumentType.CODE, 'java'),
            '.cpp': (DocumentType.CODE, 'cpp'),
            '.c': (DocumentType.CODE, 'c'),
            '.php': (DocumentType.CODE, 'php'),
            '.rb': (DocumentType.CODE, 'ruby'),
            '.go': (DocumentType.CODE, 'go'),
            
            '.xlsx': (DocumentType.DATA, 'excel'),
            '.xls': (DocumentType.DATA, 'excel'),
            '.csv': (DocumentType.DATA, 'csv'),
            '.json': (DocumentType.DATA, 'json'),
            '.xml': (DocumentType.DATA, 'xml'),
            '.sql': (DocumentType.DATA, 'sql'),
            
            '.pptx': (DocumentType.PRESENTATIONS, 'powerpoint'),
            '.ppt': (DocumentType.PRESENTATIONS, 'powerpoint'),
            '.odp': (DocumentType.PRESENTATIONS, 'libreoffice'),
            
            '.zip': (DocumentType.COMPRESSED_FILES, 'zip'),
            '.rar': (DocumentType.COMPRESSED_FILES, 'rar'),
            '.7z': (DocumentType.COMPRESSED_FILES, '7zip'),
            '.tar': (DocumentType.COMPRESSED_FILES, 'tar'),
            '.gz': (DocumentType.COMPRESSED_FILES, 'gzip'),
            
            '.exe': (DocumentType.EXECUTABLES, 'windows'),
            '.msi': (DocumentType.EXECUTABLES, 'windows_installer'),
            '.deb': (DocumentType.EXECUTABLES, 'debian'),
            '.rpm': (DocumentType.EXECUTABLES, 'redhat'),
            '.dmg': (DocumentType.EXECUTABLES, 'macos'),
        }
        
        basic_category, basic_subcategory = category_map.get(extension, (DocumentType.OTHERS, 'unknown'))
        
        if self.ai_analyzer and basic_category == DocumentType.DOCUMENTS and content:
            try:
                analysis = self.ai_analyzer.analyze_text(content, file_path)
                
                if analysis.category_suggestions:
                    suggested_category = analysis.category_suggestions[0]
                    
                    subcategory_map = {
                        'contracts': 'contracts',
                        'invoices': 'invoices', 
                        'reports': 'reports',
                        'manuals': 'manuals',
                        'certificates': 'certificates',
                        'academic': 'academic',
                        'financial': 'financial',
                        'personal': 'personal'
                    }
                    
                    if suggested_category in subcategory_map:
                        return (DocumentType.DOCUMENTS, subcategory_map[suggested_category])
                
            except Exception as e:
                logger.warning(f"Error in AI analysis for {file_path}: {e}")
        
        return (basic_category, basic_subcategory)

class DocumentOrganizer:
    def __init__(self, root_directory: Path):
        self.root_directory = Path(root_directory)
        self.input_directory = self.root_directory / "input_files"
        self.output_directory = self.root_directory / "organized_files"
        self.analyzer = FileAnalyzer()
        self.processed_files = {}
        self.stats = {
            'total_files': 0,
            'organized_files': 0,
            'skipped_files': 0,
            'errors': 0
        }
        
        self.system_files = {
            'system.log', 'processed_files.json', 'ai_analyzer.py', 'ai_gratuita.py',
            'main.py', 'config.py', 'utils.py', 'requirements.txt', 'README.md',
            '__pycache__', '.git', '.gitignore', '.venv', 'Scripts', 'Lib'
        }
        
        self._create_directories()
        self._load_processed_files()
    
    def _create_directories(self):
        self.input_directory.mkdir(exist_ok=True)
        self.output_directory.mkdir(exist_ok=True)
        
        welcome_file = self.input_directory / "README.txt"
        if not welcome_file.exists():
            with open(welcome_file, 'w', encoding='utf-8') as f:
                f.write("SMART FILE ORGANIZER - INPUT FOLDER\n")
                f.write("=====================================\n\n")
                f.write("Place your files here to be automatically organized.\n")
                f.write("The system will analyze and categorize them using AI.\n\n")
                f.write("Supported file types:\n")
                f.write("- Documents: PDF, Word, Text, etc.\n")
                f.write("- Images: JPG, PNG, GIF, etc.\n")
                f.write("- Videos: MP4, AVI, MKV, etc.\n")
                f.write("- Audio: MP3, WAV, FLAC, etc.\n")
                f.write("- Code: Python, JavaScript, HTML, etc.\n")
                f.write("- Data: Excel, CSV, JSON, etc.\n")
                f.write("- And many more...\n\n")
                f.write("Organized files will appear in the 'organized_files' folder.\n")
    
    def _load_processed_files(self):
        processed_file = self.root_directory / "processed_files.json"
        if processed_file.exists():
            try:
                with open(processed_file, 'r', encoding='utf-8') as f:
                    self.processed_files = json.load(f)
            except Exception as e:
                logger.error(f"Error loading processed files: {e}")
                self.processed_files = {}
    
    def _save_processed_files(self):
        processed_file = self.root_directory / "processed_files.json"
        try:
            with open(processed_file, 'w', encoding='utf-8') as f:
                json.dump(self.processed_files, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error saving processed files: {e}")
    
    def _is_system_file(self, file_path: Path) -> bool:
        file_name = file_path.name
        file_parent = file_path.parent.name
        
        if any(sys_file in str(file_path) for sys_file in self.system_files):
            return True
        
        if file_parent in self.system_files:
            return True
            
        if file_path.is_relative_to(self.output_directory):
            return True
            
        return False
    
    def organize_file(self, file_path: Path) -> bool:
        try:
            if self._is_system_file(file_path):
                return True
            
            file_hash = self.analyzer.get_file_hash(file_path)
            if file_hash in self.processed_files:
                logger.info(f"File already processed: {file_path}")
                self.stats['skipped_files'] += 1
                return True
            
            print(f"{Fore.CYAN}Analyzing: {file_path.name}{Style.RESET_ALL}")
            file_metadata = self.analyzer.analyze_file(file_path)
            
            if not file_metadata:
                logger.error(f"Could not analyze file: {file_path}")
                self.stats['errors'] += 1
                return False
            
            category_dir = self.output_directory / file_metadata.category.value
            subcategory_dir = category_dir / file_metadata.subcategory
            subcategory_dir.mkdir(parents=True, exist_ok=True)
            
            destination = subcategory_dir / file_metadata.name
            counter = 1
            while destination.exists():
                name_parts = file_metadata.name.rsplit('.', 1)
                if len(name_parts) == 2:
                    destination = subcategory_dir / f"{name_parts[0]}_{counter}.{name_parts[1]}"
                else:
                    destination = subcategory_dir / f"{file_metadata.name}_{counter}"
                counter += 1
            
            shutil.move(str(file_path), str(destination))
            
            self.processed_files[file_hash] = {
                'original_path': str(file_path),
                'destination_path': str(destination),
                'processed_date': datetime.now().isoformat(),
                'category': file_metadata.category.value,
                'subcategory': file_metadata.subcategory,
                'content_summary': file_metadata.content_summary
            }
            
            print(f"{Fore.GREEN}Organized: {file_path.name} -> {file_metadata.category.value}/{file_metadata.subcategory}{Style.RESET_ALL}")
            logger.info(f"File organized: {file_path} -> {destination}")
            
            self.stats['organized_files'] += 1
            return True
            
        except Exception as e:
            logger.error(f"Error organizing file {file_path}: {e}")
            print(f"{Fore.RED}Error organizing: {file_path.name} - {e}{Style.RESET_ALL}")
            self.stats['errors'] += 1
            return False
    
    def organize_directory(self, directory: Path = None) -> Dict:
        if directory is None:
            directory = self.input_directory
            
        print(f"{Fore.YELLOW}Starting automatic file organization...{Style.RESET_ALL}")
        print(f"{Fore.BLUE}Source: {directory}{Style.RESET_ALL}")
        print(f"{Fore.BLUE}Destination: {self.output_directory}{Style.RESET_ALL}")
        print("-" * 60)
        
        all_files = []
        for file_path in directory.rglob('*'):
            if file_path.is_file() and not self._is_system_file(file_path):
                all_files.append(file_path)
        
        self.stats['total_files'] = len(all_files)
        print(f"{Fore.MAGENTA}Total files found: {len(all_files)}{Style.RESET_ALL}")
        
        for i, file_path in enumerate(all_files, 1):
            print(f"\n{Fore.CYAN}[{i}/{len(all_files)}]{Style.RESET_ALL}", end=" ")
            self.organize_file(file_path)
            
            if i % 10 == 0:
                self._save_processed_files()
        
        self._save_processed_files()
        self._show_statistics()
        
        return self.stats
    
    def _show_statistics(self):
        print("\n" + "="*60)
        print(f"{Fore.GREEN}{Style.BRIGHT}ORGANIZATION STATISTICS{Style.RESET_ALL}")
        print("="*60)
        print(f"{Fore.BLUE}Total files processed: {self.stats['total_files']}{Style.RESET_ALL}")
        print(f"{Fore.GREEN}Files organized: {self.stats['organized_files']}{Style.RESET_ALL}")
        print(f"{Fore.YELLOW}Files skipped: {self.stats['skipped_files']}{Style.RESET_ALL}")
        print(f"{Fore.RED}Errors: {self.stats['errors']}{Style.RESET_ALL}")
        
        if self.stats['organized_files'] > 0:
            print(f"\n{Fore.CYAN}Structure created in: {self.output_directory}{Style.RESET_ALL}")
            self._show_directory_structure()
    
    def _show_directory_structure(self):
        print(f"\n{Fore.MAGENTA}Created folder structure:{Style.RESET_ALL}")
        for category_dir in self.output_directory.iterdir():
            if category_dir.is_dir():
                file_count = sum(1 for _ in category_dir.rglob('*') if _.is_file())
                print(f"  {category_dir.name}/ ({file_count} files)")
                for subdir in category_dir.iterdir():
                    if subdir.is_dir():
                        sub_file_count = sum(1 for _ in subdir.rglob('*') if _.is_file())
                        print(f"    {subdir.name}/ ({sub_file_count} files)")

class FileWatcher(FileSystemEventHandler):
    def __init__(self, organizer: DocumentOrganizer):
        self.organizer = organizer
        
    def on_created(self, event):
        if not event.is_directory:
            file_path = Path(event.src_path)
            if not self.organizer._is_system_file(file_path):
                time.sleep(1)
                print(f"\n{Fore.CYAN}New file detected: {file_path.name}{Style.RESET_ALL}")
                self.organizer.organize_file(file_path)

def main():
    print(f"""
{Fore.GREEN}{Style.BRIGHT}
================================================================
                                                               
              SMART FILE ORGANIZER AI SYSTEM                  
                                                               
================================================================
{Style.RESET_ALL}
    """)
    
    if AI_AVAILABLE:
        print(f"{Fore.GREEN}Free AI available{Style.RESET_ALL}")
        try:
            check_dependencies()
        except Exception as e:
            logger.warning(f"Error checking dependencies: {e}")
    else:
        print(f"{Fore.YELLOW}AI not available - will work with basic analysis{Style.RESET_ALL}")
    
    current_dir = Path.cwd()
    print(f"{Fore.BLUE}Current directory: {current_dir}{Style.RESET_ALL}")
    
    organizer = DocumentOrganizer(current_dir)
    
    while True:
        print(f"\n{Fore.YELLOW}What would you like to do?{Style.RESET_ALL}")
        print("1. Organize files from input folder")
        print("2. Monitor input folder in real-time")
        print("3. View statistics")
        print("4. View folder structure")
        print("5. Open input folder")
        print("6. Open output folder")
        print("7. Exit")
        
        choice = input(f"\n{Fore.CYAN}Select an option (1-7): {Style.RESET_ALL}").strip()
        
        if choice == "1":
            organizer.organize_directory()
            
        elif choice == "2":
            print(f"\n{Fore.GREEN}Starting real-time monitoring...{Style.RESET_ALL}")
            print(f"{Fore.YELLOW}Press Ctrl+C to stop{Style.RESET_ALL}")
            
            event_handler = FileWatcher(organizer)
            observer = Observer()
            observer.schedule(event_handler, str(organizer.input_directory), recursive=True)
            observer.start()
            
            try:
                while True:
                    time.sleep(1)
            except KeyboardInterrupt:
                observer.stop()
                print(f"\n{Fore.RED}Monitoring stopped{Style.RESET_ALL}")
            observer.join()
            
        elif choice == "3":
            organizer._show_statistics()
            
        elif choice == "4":
            organizer._show_directory_structure()
            
        elif choice == "5":
            try:
                os.startfile(str(organizer.input_directory))
                print(f"{Fore.GREEN}Input folder opened{Style.RESET_ALL}")
            except Exception as e:
                print(f"{Fore.RED}Could not open folder: {e}{Style.RESET_ALL}")
                
        elif choice == "6":
            try:
                os.startfile(str(organizer.output_directory))
                print(f"{Fore.GREEN}Output folder opened{Style.RESET_ALL}")
            except Exception as e:
                print(f"{Fore.RED}Could not open folder: {e}{Style.RESET_ALL}")
            
        elif choice == "7":
            print(f"{Fore.GREEN}Goodbye! Your files are organized.{Style.RESET_ALL}")
            break
            
        else:
            print(f"{Fore.RED}Invalid option. Please select 1-7.{Style.RESET_ALL}")

if __name__ == "__main__":
    main()
